import os
import time
import math
import re
import asyncio
from datetime import datetime
from operator import itemgetter
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import docx
from docx.shared import Inches
import numpy as np
import pandas as pd
from tqdm import tqdm

# 模型与向量库配置集中在 config.py
from config import embeddings, rerank_embeddings, llm, retriever, vectorstore

# 业务模块
from data_analysis import DataGenerator, StatisticalAnalyzer, DataProcessor
from visualization import ChartGenerator
from math_utils import OptimizationSolver, StatisticalModel, NumericalMethods, ModelEvaluator
from web_data import WebDataCache, TavilyClient, fetch_authoritative_tables_with_cache
from tavily_processing import process_tavily_payload
from iterative_generation import expand_until_length, LengthSpec
from prompts import (
    GLOBAL_AGENT_PROMPT,
    ABSTRACT_PROMPT_TEMPLATE,
    PROBLEM_RESTATEMENT_PROMPT,
    PROBLEM_ANALYSIS_PROMPT,
    MODEL_ASSUMPTION_PROMPT,
    SYMBOL_EXPLANATION_PROMPT,
    MODEL_BUILDING_PROMPT,
    MODEL_SOLVING_PROMPT,
    RESULTS_ANALYSIS_PROMPT,
)

# ===================== 全局检索缓存 =====================
# 全局缓存字典，避免重复检索相同查询
_retrieval_cache = {}
_retrieval_lock = asyncio.Lock()


def _keyword_score(query: str, text: str) -> float:
    """非常轻量的关键词匹配打分（统计关键词出现次数）。"""
    if not query or not text:
        return 0.0
    # 按中英文常见分隔符切词（简易版）
    tokens = [t.strip() for t in re.split(r"[ ,，。；;、\n\t]", query) if t.strip()]
    if not tokens:
        return 0.0
    score = 0
    for tok in tokens:
        score += text.count(tok)
    return float(score)


def _hybrid_retrieve(query: str, k: int = 5, initial_k: int = 12):
    """
    混合检索：
    1）先用 FAISS 做向量召回 initial_k 条；
    2）用关键词匹配 + 轻量 bge-reranker-v2-m3 做二次重排序；
    3）返回前 k 条。
    """
    # 1. 向量召回（带相似度分数）
    results = vectorstore.similarity_search_with_score(query, k=initial_k)
    if not results:
        return []

    docs, dist_scores = zip(*results)

    # 将距离转换为相似度（距离越小，相似度越大）
    vec_sims = [1.0 / (1.0 + float(d)) for d in dist_scores]

    # 2. 关键词得分
    kw_scores = [_keyword_score(query, d.page_content or "") for d in docs]
    if max(kw_scores) > 0:
        kw_scores = [s / max(kw_scores) for s in kw_scores]
    else:
        kw_scores = [0.0 for _ in kw_scores]

    # 3. 轻量重排序：bge-reranker-v2-m3 做打分
    texts = [d.page_content or "" for d in docs]
    q_emb = rerank_embeddings.embed_query(query)
    doc_embs = rerank_embeddings.embed_documents(texts)

    rerank_sims = []
    for emb in doc_embs:
        # 余弦相似度
        dot = sum(a * b for a, b in zip(q_emb, emb))
        q_norm = math.sqrt(sum(a * a for a in q_emb)) or 1.0
        d_norm = math.sqrt(sum(a * a for a in emb)) or 1.0
        rerank_sims.append(dot / (q_norm * d_norm))

    # 归一化到 0-1
    min_r, max_r = min(rerank_sims), max(rerank_sims)
    if max_r > min_r:
        rerank_sims = [(s - min_r) / (max_r - min_r) for s in rerank_sims]
    else:
        rerank_sims = [0.5 for _ in rerank_sims]

    # 4. 融合三种得分：向量相似 + 关键词 + reranker
    final_scores = []
    for v, k_s, r in zip(vec_sims, kw_scores, rerank_sims):
        # 权重可以按需调整：reranker 为主，向量 + 关键词 辅助
        score = 0.6 * r + 0.3 * v + 0.1 * k_s
        final_scores.append(score)

    ranked = sorted(zip(docs, final_scores), key=lambda x: x[1], reverse=True)
    top_docs = [d for d, _ in ranked[:k]]
    return top_docs


def get_cached_retrieval(query: str):
    """
    带缓存的检索函数
    :param query: 检索查询字符串
    :return: 格式化的检索上下文
    """
    if query not in _retrieval_cache:
        docs = _hybrid_retrieve(query, k=5, initial_k=12)
        context = format_docs(docs)
        _retrieval_cache[query] = (docs, context)
        return docs, context
    return _retrieval_cache[query]


async def get_cached_retrieval_async(query: str):
    """
    异步版缓存检索：使用锁避免并发重复检索同一 query。
    注意：FAISS 检索是同步的，这里用 asyncio.to_thread 放到线程池执行。
    """
    async with _retrieval_lock:
        if query in _retrieval_cache:
            return _retrieval_cache[query]

    # 不在锁内执行耗时操作，避免阻塞其他 key 的读取
    docs = await asyncio.to_thread(_hybrid_retrieve, query, 5, 12)
    context = format_docs(docs)

    async with _retrieval_lock:
        _retrieval_cache[query] = (docs, context)
        return _retrieval_cache[query]

# ===================== 3. 核心生成函数（LCEL链式调用） =====================
def format_docs(docs):
    """格式化FAISS检索的知识库内容"""
    # 控制检索上下文长度，避免 prompt 过长导致 Ollama 推理变慢
    max_doc_chars = 800
    max_total_chars = 3200
    parts = []
    total = 0
    for i, doc in enumerate(docs):
        content = (doc.page_content or "").strip().replace("\u0000", "")
        if len(content) > max_doc_chars:
            content = content[:max_doc_chars] + "…【已截断】"
        piece = f"参考资料{i + 1}：{content}"
        if total + len(piece) > max_total_chars:
            break
        parts.append(piece)
        total += len(piece)
    return "\n\n".join(parts)


# 全局变量，用于存储进度回调函数
_global_progress_callback = None

# 最近一次 Tavily 爬取结果（供 API/前端读取）
_last_tavily_payload = None
_last_tavily_run_dir = None


def get_last_tavily_payload():
    return _last_tavily_payload


def get_last_tavily_run_dir():
    return _last_tavily_run_dir


def _load_tavily_constraints() -> str:
    path = os.path.join(os.path.dirname(__file__), "tavily_query_constraints.txt")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception:
        return ""


def _new_tavily_run_dir() -> str:
    base = os.path.join(os.path.dirname(__file__), "tavily_runs")
    os.makedirs(base, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    run_dir = os.path.join(base, ts)
    os.makedirs(run_dir, exist_ok=True)
    return run_dir

def set_progress_callback(callback):
    """设置全局进度回调函数"""
    global _global_progress_callback
    _global_progress_callback = callback

def _timed(label: str):
    """简单计时器：打印每一步耗时，方便定位卡顿点。"""
    start = time.time()

    def done(extra: str = ""):
        cost = time.time() - start
        tail = f" {extra}" if extra else ""
        message = f"{label} 耗时：{cost:.1f}s{tail}"
        tqdm.write(f"    - {message}")
        # 通过全局回调函数传递耗时信息
        if _global_progress_callback:
            _global_progress_callback(None, message)

    return done


def generate_section(chain, input_data):
    """通用生成函数：调用LCEL链生成指定章节内容"""
    return chain.invoke(input_data).strip()


# 3.1 生成摘要
def generate_abstract(user_prompt):
    # 检索知识库相关内容（使用缓存）
    t = _timed("摘要-检索")
    docs, context = get_cached_retrieval(user_prompt)
    t(f"(命中{len(docs)}条)")
    # 构建LCEL链（检索上下文 + 用户提示词 → 摘要生成）
    abstract_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"]
            )
            | ABSTRACT_PROMPT_TEMPLATE
            | llm
            | StrOutputParser()
    )
    t2 = _timed("摘要-生成")
    out = generate_section(abstract_chain, {"user_prompt": user_prompt})
    t2(f"(输出{len(out)}字)")
    return out


# 3.2 生成问题重述
def generate_problem_restatement(user_prompt, abstract_content):
    t = _timed("问题重述-检索")
    docs, context = get_cached_retrieval(user_prompt + "问题重述")
    t(f"(命中{len(docs)}条)")
    restatement_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"]
            )
            | PROBLEM_RESTATEMENT_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("问题重述-生成")
    out = generate_section(restatement_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content
    })
    t2(f"(输出{len(out)}字)")
    return out


# 3.3 生成问题分析
def generate_problem_analysis(user_prompt, abstract_content, problem_restatement_content):
    t = _timed("问题分析-检索")
    docs, context = get_cached_retrieval(user_prompt + "问题分析")
    t(f"(命中{len(docs)}条)")
    analysis_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_restatement_content=lambda x: x["problem_restatement_content"]
            )
            | PROBLEM_ANALYSIS_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("问题分析-生成")
    out = generate_section(analysis_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_restatement_content": problem_restatement_content
    })
    t2(f"(输出{len(out)}字)")
    return out


# 3.4 生成模型假设
def generate_model_assumption(user_prompt, abstract_content, problem_analysis_content):
    t = _timed("模型假设-检索")
    docs, context = get_cached_retrieval(user_prompt + "模型假设")
    t(f"(命中{len(docs)}条)")
    assumption_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_analysis_content=lambda x: x["problem_analysis_content"]
            )
            | MODEL_ASSUMPTION_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("模型假设-生成")
    out = generate_section(assumption_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_analysis_content": problem_analysis_content
    })
    t2(f"(输出{len(out)}字)")
    return out


# 3.5 生成符号说明
def generate_symbol_explanation(user_prompt, abstract_content, problem_analysis_content):
    t = _timed("符号说明-检索")
    docs, context = get_cached_retrieval(user_prompt + "符号说明")
    t(f"(命中{len(docs)}条)")
    symbol_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_analysis_content=lambda x: x["problem_analysis_content"]
            )
            | SYMBOL_EXPLANATION_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("符号说明-生成")
    out = generate_section(symbol_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_analysis_content": problem_analysis_content
    })
    t2(f"(输出{len(out)}字)")
    return out


# 3.6 生成模型建立（新增）
def generate_model_building(user_prompt, abstract_content, problem_analysis_content,
                           model_assumption_content, symbol_explanation_content, data_analysis_results,
                           conversation_context: str = ""):
    docs, context = get_cached_retrieval(user_prompt + "模型建立")
    building_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_analysis_content=lambda x: x["problem_analysis_content"],
                model_assumption_content=lambda x: x["model_assumption_content"],
                symbol_explanation_content=lambda x: x["symbol_explanation_content"],
                data_analysis_results=lambda x: x["data_analysis_results"]
            )
            | MODEL_BUILDING_PROMPT
            | llm
            | StrOutputParser()
    )
    draft = generate_section(building_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_analysis_content": problem_analysis_content,
        "model_assumption_content": model_assumption_content,
        "symbol_explanation_content": symbol_explanation_content,
        "data_analysis_results": data_analysis_results
    })
    # 多轮扩写：字符数目标近似对应 prompts 中的“字数”
    return expand_until_length(
        llm=llm,
        draft=draft,
        length_spec=LengthSpec(target=2000, tolerance_ratio=0.10),
        requirements="模型建立：严格按3.1/3.2/3.3/3.4分段；公式LaTeX+编号；步骤可复现；符号不冲突；禁用模糊措辞。",
        context_blob=f"【赛题】{user_prompt}\n\n【多轮对话补充】{conversation_context}\n\n【知识库】{context}\n\n【摘要】{abstract_content}\n\n【问题分析】{problem_analysis_content}\n\n【模型假设】{model_assumption_content}\n\n【符号说明】{symbol_explanation_content}\n\n【数据分析结论/外部表格】{data_analysis_results}",
        max_rounds=4,
        progress_cb=lambda msg: (_global_progress_callback(None, msg) if _global_progress_callback else None),
    )


# 3.7 生成模型求解（新增）
def generate_model_solving(user_prompt, model_building_content, solving_results, conversation_context: str = ""):
    docs, context = get_cached_retrieval(user_prompt + "模型求解")
    solving_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                model_building_content=lambda x: x["model_building_content"],
                solving_results=lambda x: x["solving_results"]
            )
            | MODEL_SOLVING_PROMPT
            | llm
            | StrOutputParser()
    )
    draft = generate_section(solving_chain, {
        "user_prompt": user_prompt,
        "model_building_content": model_building_content,
        "solving_results": solving_results
    })
    return expand_until_length(
        llm=llm,
        draft=draft,
        length_spec=LengthSpec(target=5000, tolerance_ratio=0.10),
        requirements="模型求解：4.1/4.2/4.3分段；算法步骤可复现；参数具体到数值；结果含单位与关键数值；工具/库/函数明确。",
        context_blob=f"【赛题】{user_prompt}\n\n【多轮对话补充】{conversation_context}\n\n【知识库】{context}\n\n【模型建立】{model_building_content}\n\n【求解结果提示】{solving_results}",
        max_rounds=4,
        progress_cb=lambda msg: (_global_progress_callback(None, msg) if _global_progress_callback else None),
    )


# 3.8 生成结果分析（新增）
def generate_results_analysis(user_prompt, model_solving_content, chart_files, statistical_results, conversation_context: str = ""):
    docs, context = get_cached_retrieval(user_prompt + "结果分析")
    analysis_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                model_solving_content=lambda x: x["model_solving_content"],
                chart_files=lambda x: x["chart_files"],
                statistical_results=lambda x: x["statistical_results"]
            )
            | RESULTS_ANALYSIS_PROMPT
            | llm
            | StrOutputParser()
    )
    draft = generate_section(analysis_chain, {
        "user_prompt": user_prompt,
        "model_solving_content": model_solving_content,
        "chart_files": chart_files,
        "statistical_results": statistical_results
    })
    return expand_until_length(
        llm=llm,
        draft=draft,
        length_spec=LengthSpec(target=4500, tolerance_ratio=0.10),
        requirements="结果分析：5.1/5.2/5.3分段；必须结合图表清单逐条引用并解读；包含敏感性/误差分析且给具体数值；结合赛题现实意义。",
        context_blob=f"【赛题】{user_prompt}\n\n【多轮对话补充】{conversation_context}\n\n【知识库】{context}\n\n【模型求解】{model_solving_content}\n\n【图表清单】{chart_files}\n\n【统计分析结果】{statistical_results}",
        max_rounds=4,
        progress_cb=lambda msg: (_global_progress_callback(None, msg) if _global_progress_callback else None),
    )


# ===================== 异步生成（用于并行非依赖章节） =====================
async def generate_model_assumption_async(user_prompt, abstract_content, problem_analysis_content):
    t = _timed("模型假设-检索(异步)")
    docs, context = await get_cached_retrieval_async(user_prompt + "模型假设")
    t(f"(命中{len(docs)}条)")
    assumption_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_analysis_content=lambda x: x["problem_analysis_content"]
            )
            | MODEL_ASSUMPTION_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("模型假设-生成(异步)")
    out = await asyncio.to_thread(generate_section, assumption_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_analysis_content": problem_analysis_content
    })
    t2(f"(输出{len(out)}字)")
    return out


async def generate_symbol_explanation_async(user_prompt, abstract_content, problem_analysis_content):
    t = _timed("符号说明-检索(异步)")
    docs, context = await get_cached_retrieval_async(user_prompt + "符号说明")
    t(f"(命中{len(docs)}条)")
    symbol_chain = (
            RunnablePassthrough.assign(
                context=lambda x: context,
                user_prompt=lambda x: x["user_prompt"],
                abstract_content=lambda x: x["abstract_content"],
                problem_analysis_content=lambda x: x["problem_analysis_content"]
            )
            | SYMBOL_EXPLANATION_PROMPT
            | llm
            | StrOutputParser()
    )
    t2 = _timed("符号说明-生成(异步)")
    out = await asyncio.to_thread(generate_section, symbol_chain, {
        "user_prompt": user_prompt,
        "abstract_content": abstract_content,
        "problem_analysis_content": problem_analysis_content
    })
    t2(f"(输出{len(out)}字)")
    return out


# 3.9 数据分析和可视化生成函数（新增）
def generate_data_and_charts(user_prompt, problem_type="general"):
    """
    根据问题类型生成示例数据和图表
    :param user_prompt: 用户提示词
    :param problem_type: 问题类型 ('time_series', 'optimization', 'regression', 'general')
    :return: (数据字典, 图表文件列表, 统计分析结果)
    """
    tqdm.write("  生成示例数据和图表...")
    chart_generator = ChartGenerator()
    data_generator = DataGenerator()
    analyzer = StatisticalAnalyzer()
    
    generated_data = {}
    chart_files = []
    statistical_results = {}
    
    # 根据问题类型生成不同类型的数据和图表
    if "时间" in user_prompt or "预测" in user_prompt or "序列" in user_prompt:
        # 时间序列数据
        time_data = data_generator.generate_time_series_data(n_points=100, trend="linear", seasonality=True)
        generated_data['time_series'] = time_data
        
        # 生成时间序列图
        chart_file = chart_generator.plot_time_series(time_data, title="时间序列数据可视化")
        chart_files.append(("图1：时间序列数据可视化", chart_file))
        
        # 统计分析
        stats_result = analyzer.time_series_analysis(time_data)
        statistical_results['time_series'] = stats_result
    
    if "优化" in user_prompt or "分配" in user_prompt or "规划" in user_prompt:
        # 优化问题数据
        opt_data = data_generator.generate_optimization_data(n_items=20)
        generated_data['optimization'] = opt_data
        
        # 生成柱状图
        chart_file = chart_generator.plot_bar_chart(
            opt_data.head(10), '物品编号', '价值', title="物品价值分布"
        )
        chart_files.append(("图2：物品价值分布", chart_file))
    
    if "回归" in user_prompt or "关系" in user_prompt or "影响" in user_prompt:
        # 多变量数据
        multi_data = data_generator.generate_multivariate_data(n_samples=200, n_features=5)
        generated_data['multivariate'] = multi_data
        
        # 生成散点图
        chart_file = chart_generator.plot_scatter(
            multi_data, '特征1', '目标变量', title="特征与目标变量关系"
        )
        chart_files.append(("图3：特征与目标变量关系", chart_file))
        
        # 生成相关性热力图
        chart_file = chart_generator.plot_correlation_heatmap(multi_data, title="变量相关性热力图")
        chart_files.append(("图4：变量相关性热力图", chart_file))
        
        # 回归分析
        reg_result = analyzer.regression_analysis(multi_data, '特征1', '目标变量')
        statistical_results['regression'] = reg_result
    
    # 通用统计图表
    if len(generated_data) > 0:
        # 选择第一个数据集生成直方图
        first_data = list(generated_data.values())[0]
        numeric_cols = first_data.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            chart_file = chart_generator.plot_histogram(
                first_data, numeric_cols[0], title="数据分布直方图"
            )
            chart_files.append(("图5：数据分布直方图", chart_file))
            
            # 基本统计量
            stats_result = analyzer.basic_statistics(first_data, numeric_cols[0])
            statistical_results['basic_stats'] = stats_result
    
    # 格式化结果用于提示词
    chart_files_str = "\n".join([f"{name}: {file}" for name, file in chart_files])
    statistical_results_str = "\n".join([
        f"{key}: {value}" for key, value in statistical_results.items()
    ])
    
    return generated_data, chart_files, statistical_results_str


# ===================== 4. 整合生成Word文档 =====================
def generate_math_modeling_paper(user_prompt, save_path="数学建模论文.docx",
                                generate_charts=True, generate_data=True,
                                progress_callback=None,
                                conversation_context: str = ""):
    """
    生成完整的数学建模论文Word文档（增强版，包含数据分析和图表）
    :param user_prompt: 用户提供的赛题描述（几百字的长提示词）
    :param save_path: 论文保存路径
    :param generate_charts: 是否生成图表
    :param generate_data: 是否生成示例数据
    :param progress_callback: 进度回调函数，接收(progress, message)参数
    """
    def update_progress(progress, message):
        """更新进度"""
        if progress_callback:
            progress_callback(progress, message)
    
    # 设置全局进度回调函数，以便 _timed 函数能够传递耗时信息
    set_progress_callback(progress_callback)
    
    tqdm.write("=== 开始生成数学建模论文（增强版） ===")
    
    # 定义所有步骤
    steps = [
        ("生成示例数据和图表", generate_data or generate_charts),
        ("联网抓取权威数据表格（可选）", True),
        ("生成摘要", True),
        ("生成问题重述", True),
        ("生成问题分析", True),
        ("生成模型假设", True),
        ("生成符号说明", True),
        ("生成模型建立", True),
        ("生成模型求解", True),
        ("生成结果分析", True),
        ("整合Word文档", True),
    ]
    
    # 创建总进度条
    total_steps = sum(1 for _, enabled in steps if enabled)
    pbar = tqdm(total=total_steps, desc="论文生成进度", unit="步骤", 
                bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]')
    
    # 步骤0：生成示例数据和图表（如果需要）
    generated_data = {}
    chart_files = []
    statistical_results = ""
    data_analysis_results = ""
    web_data_summary = ""
    web_data_citations = ""
    tavily_analysis_text = ""
    tavily_charts = []
    tavily_raw_csvs = []
    tavily_cleaned_csvs = []
    
    if generate_data or generate_charts:
        update_progress(10, "正在生成示例数据和图表...（步骤1/11）")
        pbar.set_description("生成示例数据和图表")
        generated_data, chart_files, statistical_results = generate_data_and_charts(user_prompt)
        data_analysis_results = f"已生成{len(generated_data)}组示例数据，{len(chart_files)}个图表"
        pbar.update(1)

    # 步骤0.5：联网抓取权威表格（可选：有 TAVILY_API_KEY 才会启用）
    update_progress(20, "正在联网抓取权威数据表格...（步骤2/11）")
    pbar.set_description("联网抓取权威数据表格（可选）")
    try:
        tavily = TavilyClient()
        if tavily.enabled():
            cache = WebDataCache()
            # 针对建模题常见的“数据来源”检索模板：你也可以按赛题再细化关键词
            query = f"{user_prompt} 权威 数据 表格 统计 年鉴 官方"
            # 为本次运行创建 tavily_runs 子目录（按时间戳分批次）
            tavily_base_dir = os.path.join(os.path.dirname(__file__), "tavily_runs")
            os.makedirs(tavily_base_dir, exist_ok=True)
            latest_dir = _new_tavily_run_dir()

            payload = fetch_authoritative_tables_with_cache(
                query, cache, tavily, max_sources=3, max_tables_per_source=2, run_dir=latest_dir
            )
            web_data_summary = payload.get("summary", "")
            # 生成引用列表（最多 3 条）
            sources = payload.get("sources", [])[:3]
            web_data_citations = "\n".join([f"- {s.get('title','')}：{s.get('url','')}" for s in sources])

            # 对 tavily 表格做通用清洗 + 分析 + 可视化
            try:
                tavily_analysis_text, tavily_charts, tavily_cleaned_csvs, tavily_raw_csvs = process_tavily_payload(
                    payload, os.path.join(latest_dir, "processed")
                )
                # 记录供 API/前端读取
                global _last_tavily_payload, _last_tavily_run_dir
                _last_tavily_payload = payload
                _last_tavily_run_dir = latest_dir
            except Exception as e:
                tavily_analysis_text = f"外部权威表格自动分析失败：{e}"
                tavily_charts = []
                tavily_raw_csvs = []
                tavily_cleaned_csvs = []
        else:
            web_data_summary = "未启用联网数据抓取（未检测到环境变量 TAVILY_API_KEY）。"
            web_data_citations = ""
    except Exception as e:
        web_data_summary = f"联网数据抓取失败：{e}"
        web_data_citations = ""
    pbar.update(1)

    # 将 Tavily 的结构化分析文本并入数据分析结论，供后续“模型建立/结果分析”引用扩写
    if tavily_analysis_text:
        data_analysis_results = (data_analysis_results + "\n\n【外部权威表格结构化分析】\n" + tavily_analysis_text).strip()
    
    # 步骤1：生成摘要
    update_progress(30, "正在生成摘要...（步骤3/11）")
    pbar.set_description("生成摘要")
    abstract = generate_abstract(user_prompt + ("\n\n【外部权威数据摘要】\n" + web_data_summary if web_data_summary else ""))
    pbar.update(1)

    # 步骤2：生成问题重述
    update_progress(40, "正在生成问题重述...（步骤4/11）")
    pbar.set_description("生成问题重述")
    problem_restatement = generate_problem_restatement(user_prompt, abstract)
    pbar.update(1)

    # 步骤3：生成问题分析
    update_progress(50, "正在生成问题分析...（步骤5/11）")
    pbar.set_description("生成问题分析")
    problem_analysis = generate_problem_analysis(user_prompt, abstract, problem_restatement)
    pbar.update(1)

    # 步骤4/5：并行生成（非依赖型章节）
    # - 模型假设、符号说明只依赖摘要/问题分析，因此可并行以减少整体耗时
    update_progress(60, "正在并行生成模型假设和符号说明...（步骤6-7/11）")
    pbar.set_description("并行生成：模型假设 & 符号说明")

    async def _run_parallel():
        return await asyncio.gather(
            generate_model_assumption_async(user_prompt, abstract, problem_analysis),
            generate_symbol_explanation_async(user_prompt, abstract, problem_analysis),
        )

    model_assumption, symbol_explanation = asyncio.run(_run_parallel())
    # 两个步骤都完成，进度条更新两次
    pbar.update(2)

    # 步骤6：生成模型建立（新增）
    update_progress(70, "正在生成模型建立...（步骤8/11）")
    pbar.set_description("生成模型建立")
    model_building = generate_model_building(
        user_prompt, abstract, problem_analysis, 
        model_assumption, symbol_explanation,
        data_analysis_results + ("\n\n【外部权威数据摘要】\n" + web_data_summary if web_data_summary else "") + ("\n\n【参考链接】\n" + web_data_citations if web_data_citations else ""),
        conversation_context=conversation_context
    )
    pbar.update(1)
    
    # 步骤7：生成模型求解（新增）
    update_progress(80, "正在生成模型求解...（步骤9/11）")
    pbar.set_description("生成模型求解")
    solving_results = f"使用Python和NumPy/SciPy进行数值计算，已生成{len(chart_files)}个可视化结果"
    model_solving = generate_model_solving(user_prompt, model_building, solving_results, conversation_context=conversation_context)
    pbar.update(1)
    
    # 步骤8：生成结果分析（新增）
    update_progress(90, "正在生成结果分析...（步骤10/11）")
    pbar.set_description("生成结果分析")
    chart_files_str = "\n".join([f"{name}" for name, _ in chart_files])
    results_analysis = generate_results_analysis(
        user_prompt, model_solving, chart_files_str, statistical_results, conversation_context=conversation_context
    )
    pbar.update(1)

    # 步骤9：整合到Word文档
    update_progress(95, "正在整合Word文档...（步骤11/11）")
    pbar.set_description("整合Word文档")
    doc = docx.Document()

    # 添加标题（从摘要中提取，或自主生成）
    title = abstract.split("\n")[0].replace("**", "").strip() if "**" in abstract else "数学建模论文"
    doc.add_heading(title, level=0)

    # 添加摘要
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(abstract)

    # 添加问题重述
    doc.add_heading("问题重述", level=1)
    doc.add_paragraph(problem_restatement)

    # 添加问题分析
    doc.add_heading("问题分析", level=1)
    doc.add_paragraph(problem_analysis)

    # 添加模型假设
    doc.add_heading("模型假设", level=1)
    doc.add_paragraph(model_assumption)

    # 添加符号说明
    doc.add_heading("符号说明", level=1)
    doc.add_paragraph(symbol_explanation)
    
    # 添加模型建立（新增）
    doc.add_heading("模型建立", level=1)
    doc.add_paragraph(model_building)
    
    # 添加模型求解（新增）
    doc.add_heading("模型求解", level=1)
    doc.add_paragraph(model_solving)
    
    # 添加结果分析（新增）
    doc.add_heading("结果分析", level=1)
    doc.add_paragraph(results_analysis)
    
    # 插入图表（新增）
    # 将 Tavily 可视化图表也纳入图表附录
    if tavily_charts:
        chart_files.extend(tavily_charts)

    if generate_charts and len(chart_files) > 0:
        doc.add_heading("图表附录", level=1)
        for i, (chart_name, chart_file) in enumerate(chart_files, 1):
            if os.path.exists(chart_file):
                doc.add_paragraph(f"{chart_name}")
                doc.add_picture(chart_file, width=Inches(6))  # 设置图片宽度为6英寸
                doc.add_paragraph("")  # 添加空行

    # 保存文档
    doc.save(save_path)
    update_progress(100, f"论文生成完成！保存路径：{os.path.abspath(save_path)}（共生成 {len(chart_files)} 个图表）")
    pbar.update(1)
    pbar.close()
    
    tqdm.write(f"=== 论文生成完成！保存路径：{os.path.abspath(save_path)} ===")
    tqdm.write(f"=== 共生成 {len(chart_files)} 个图表 ===")
    return save_path


# ===================== 5. 运行入口 =====================
if __name__ == "__main__":
    # 简化命令行：直接提示用户输入赛题，若留空则使用示例；输出路径固定为“数学建模论文.docx”
    print("请输入赛题描述（直接粘贴，回车结束；留空则使用示例赛题）：")
    user_prompt = input().strip()

    if not user_prompt:
        user_prompt = """
NIPT（Non-invasive Prenatal Test，即无创产前检测）是一种通过采集母体血液、检测胎儿的游离
DNA片段、分析胎儿染色体是否存在异常的产前检测技术，目的是通过早期检测确定胎儿的健康状况。
根据临床经验，畸型胎儿主要有唐氏综合征、爱德华氏综合征和帕陶氏综合征，这三种体征分别由胎儿
21 号、18号和13号“染色体游离DNA片段的比例”（简称“染色体浓度”）是否异常决定。NIPT的
准确性主要由胎儿性染色体（男胎XY，女胎XX）浓度判断。通常孕妇的孕期在10周~25周之间可以
检测胎儿性染色体浓度，且如果男胎的Y染色体浓度达到或高于4%、女胎的X染色体浓度没有异常，
则可认为NIPT的结果是基本准确的，否则难以保证结果准确性要求。同时，实际中应尽早发现不健康
的胎儿，否则会带来治疗窗口期缩短的风险，早期发现（12周以内）风险较低；中期发现（13－27周）
风险高；晚期发现（28周以后）风险极高。 
实践表明，男胎Y染色体浓度与孕妇孕周数及其身体质量指数（BMI）紧密相关。通常根据孕妇的
BMI 值进行分组（例如：[20,28)，[28,32)，[32,36)，[36,40)，40 以上）分别确定 NIPT 的时点（相对孕
期的时间点）。由于每个孕妇的年龄、BMI、孕情等存在个体差异，对所有孕妇采用简单的经验分组和
统一的检测时点进行NIPT，会对其准确性产生较大影响。因此，依据BMI对孕妇进行合理分组，确定
各不同群组的最佳NIPT时点，可以减少某些孕妇因胎儿不健康而缩短治疗窗口期所带来的潜在风险。 
为了研究各类孕妇群体合适的NIPT时点，并对检测的准确性进行分析，附件给出了某地区（大多
为高BMI）孕妇的NIPT数据。在实际检测中，经常会出现测序失败（比如：检测时点过早和不确定因
素影响等）的情况。同时为了增加检测结果的可靠性，对某些孕妇有多次采血多次检测或一次采血多次
检测的情况。试利用附件提供的数据建立数学模型研究如下问题： 
问题1  试分析胎儿Y染色体浓度与孕妇的孕周数和BMI等指标的相关特性，给出相应的关系模
型，并检验其显著性。 
问题2  临床证明，男胎孕妇的BMI是影响胎儿Y染色体浓度的最早达标时间（即浓度达到或超
过4%的最早时间）的主要因素。试对男胎孕妇的BMI进行合理分组，给出每组的BMI区间和最佳NIPT
时点，使得孕妇可能的潜在风险最小，并分析检测误差对结果的影响。 
问题3  男胎Y染色体浓度达标时间受多种因素(身高、体重、年龄等)的影响，试综合考虑这些因
素、检测误差和胎儿的Y染色体浓度达标比例（即浓度达到或超过4%的比例），根据男胎孕妇的BMI，
给出合理分组以及每组的最佳NIPT时点，使得孕妇潜在风险最小，并分析检测误差对结果的影响。 
问题4  由于孕妇和女胎都不携带Y染色体，重要的是如何判定女胎是否异常。试以女胎孕妇的21
号、18号和13号染色体非整倍体（AB列）为判定结果，综合考虑X染色体及上述染色体的Z值、GC
含量、读段数及相关比例、BMI等因素，给出女胎异常的判定方法。

""".strip()
        print("已使用示例赛题。")

    conversation_context = ""
    yn = input("是否启用多轮对话补充信息？(y/n，默认n)：").strip().lower()
    if yn.startswith("y"):
        conversation_context = input("请输入多轮对话补充信息（可留空）：").strip()

    generate_math_modeling_paper(
        user_prompt,
        save_path="数学建模论文.docx",
        conversation_context=conversation_context,
    )