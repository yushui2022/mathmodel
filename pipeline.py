import asyncio
import os
import docx
from docx.shared import Inches
from tqdm import tqdm

from generators import (
    generate_abstract,
    generate_problem_restatement,
    generate_problem_analysis,
    generate_model_assumption,
    generate_symbol_explanation,
    generate_model_assumption_async,
    generate_symbol_explanation_async,
    generate_model_building,
    generate_model_solving,
    generate_results_analysis,
    generate_data_and_charts,
)
from web_data import WebDataCache, TavilyClient, fetch_authoritative_tables_with_cache


def generate_math_modeling_paper(user_prompt, save_path="数学建模论文.docx", generate_charts=True, generate_data=True):
    tqdm.write("=== 开始生成数学建模论文（增强版） ===")

    steps = [
        ("生成示例数据和图表", generate_data or generate_charts),
        ("联网抓取权威数据表格（可选）", True),
        ("生成摘要", True),
        ("生成问题重述", True),
        ("生成问题分析", True),
        ("生成模型假设", True),
        ("生成符号说明", True),
        ("生成模型建立", True),
        ("生成模型求解", True),
        ("生成结果分析", True),
        ("整合Word文档", True),
    ]

    total_steps = sum(1 for _, enabled in steps if enabled)
    pbar = tqdm(
        total=total_steps,
        desc="论文生成进度",
        unit="步骤",
        bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]",
    )

    generated_data = {}
    chart_files = []
    statistical_results = ""
    data_analysis_results = ""
    web_data_summary = ""
    web_data_citations = ""

    if generate_data or generate_charts:
        pbar.set_description("生成示例数据和图表")
        generated_data, chart_files, statistical_results = generate_data_and_charts(user_prompt)
        data_analysis_results = f"已生成{len(generated_data)}组示例数据，{len(chart_files)}个图表"
        pbar.update(1)

    pbar.set_description("联网抓取权威数据表格（可选）")
    try:
        tavily = TavilyClient()
        if tavily.enabled():
            cache = WebDataCache()
            query = f"{user_prompt} 权威 数据 表格 统计 年鉴 官方"
            payload = fetch_authoritative_tables_with_cache(query, cache, tavily, max_sources=3, max_tables_per_source=2)
            web_data_summary = payload.get("summary", "")
            sources = payload.get("sources", [])[:3]
            web_data_citations = "\n".join([f"- {s.get('title','')}：{s.get('url','')}" for s in sources])
        else:
            web_data_summary = "未启用联网数据抓取（未检测到环境变量 TAVILY_API_KEY）。"
            web_data_citations = ""
    except Exception as e:
        web_data_summary = f"联网数据抓取失败：{e}"
        web_data_citations = ""
    pbar.update(1)

    pbar.set_description("生成摘要")
    abstract = generate_abstract(user_prompt + ("\n\n【外部权威数据摘要】\n" + web_data_summary if web_data_summary else ""))
    pbar.update(1)

    pbar.set_description("生成问题重述")
    problem_restatement = generate_problem_restatement(user_prompt, abstract)
    pbar.update(1)

    pbar.set_description("生成问题分析")
    problem_analysis = generate_problem_analysis(user_prompt, abstract, problem_restatement)
    pbar.update(1)

    pbar.set_description("并行生成：模型假设 & 符号说明")

    async def _run_parallel():
        return await asyncio.gather(
            generate_model_assumption_async(user_prompt, abstract, problem_analysis),
            generate_symbol_explanation_async(user_prompt, abstract, problem_analysis),
        )

    model_assumption, symbol_explanation = asyncio.run(_run_parallel())
    pbar.update(2)

    pbar.set_description("生成模型建立")
    model_building = generate_model_building(
        user_prompt,
        abstract,
        problem_analysis,
        model_assumption,
        symbol_explanation,
        data_analysis_results
        + ("\n\n【外部权威数据摘要】\n" + web_data_summary if web_data_summary else "")
        + ("\n\n【参考链接】\n" + web_data_citations if web_data_citations else ""),
    )
    pbar.update(1)

    pbar.set_description("生成模型求解")
    solving_results = f"使用Python和NumPy/SciPy进行数值计算，已生成{len(chart_files)}个可视化结果"
    model_solving = generate_model_solving(user_prompt, model_building, solving_results)
    pbar.update(1)

    pbar.set_description("生成结果分析")
    chart_files_str = "\n".join([f"{name}" for name, _ in chart_files])
    results_analysis = generate_results_analysis(user_prompt, model_solving, chart_files_str, statistical_results)
    pbar.update(1)

    pbar.set_description("整合Word文档")
    doc = docx.Document()
    title = abstract.split("\n")[0].replace("**", "").strip() if "**" in abstract else "数学建模论文"
    doc.add_heading(title, level=0)
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(abstract)
    doc.add_heading("问题重述", level=1)
    doc.add_paragraph(problem_restatement)
    doc.add_heading("问题分析", level=1)
    doc.add_paragraph(problem_analysis)
    doc.add_heading("模型假设", level=1)
    doc.add_paragraph(model_assumption)
    doc.add_heading("符号说明", level=1)
    doc.add_paragraph(symbol_explanation)
    doc.add_heading("模型建立", level=1)
    doc.add_paragraph(model_building)
    doc.add_heading("模型求解", level=1)
    doc.add_paragraph(model_solving)
    doc.add_heading("结果分析", level=1)
    doc.add_paragraph(results_analysis)

    if generate_charts and len(chart_files) > 0:
        doc.add_heading("图表附录", level=1)
        for _, chart_file in chart_files:
            if os.path.exists(chart_file):
                doc.add_picture(chart_file, width=Inches(6))
                doc.add_paragraph("")

    doc.save(save_path)
    pbar.update(1)
    pbar.close()

    tqdm.write(f"=== 论文生成完成！保存路径：{os.path.abspath(save_path)} ===")
    tqdm.write(f"=== 共生成 {len(chart_files)} 个图表 ===")
    return save_path


__all__ = ["generate_math_modeling_paper"]

